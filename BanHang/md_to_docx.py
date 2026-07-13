#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown Báo Cáo to DOCX (Word format)
Converts BaoCao_ThucTap_Complete_Merged.md to DOCX
"""

import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx not installed")
    print("📥 Install with: pip install python-docx")
    exit(1)

# Paths
DOCS_DIR = Path("/home/nvs1512/Project IT/San-tim-vien/BanHang/docs")
MD_FILE = DOCS_DIR / "BaoCao_ThucTap_Complete_Merged.md"
DOCX_FILE = DOCS_DIR / "BaoCao_ThucTap_Complete_Merged.docx"


def add_heading_style(doc, text, level):
    """Add heading with appropriate style"""
    if level == 1:
        return doc.add_heading(text, level=1)
    elif level == 2:
        return doc.add_heading(text, level=2)
    elif level == 3:
        return doc.add_heading(text, level=3)
    else:
        return doc.add_heading(text, level=level)


def format_code_block(doc, code_text, language=""):
    """Add formatted code block to document"""
    p = doc.add_paragraph()
    p.style = 'List Bullet'

    # Add code with monospace font
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Background color (light gray)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E8E8E8')
    p._element.get_or_add_pPr().append(shading_elm)


def parse_and_convert(md_content):
    """Parse markdown and convert to DOCX"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('# '):
            heading_text = line.replace('# ', '').strip()
            add_heading_style(doc, heading_text, 1)
            i += 1
        elif line.startswith('## '):
            heading_text = line.replace('## ', '').strip()
            add_heading_style(doc, heading_text, 2)
            i += 1
        elif line.startswith('### '):
            heading_text = line.replace('### ', '').strip()
            add_heading_style(doc, heading_text, 3)
            i += 1
        elif line.startswith('#### '):
            heading_text = line.replace('#### ', '').strip()
            add_heading_style(doc, heading_text, 4)
            i += 1

        # Code blocks
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1

            code_text = '\n'.join(code_lines).strip()
            if code_text:
                format_code_block(doc, code_text)
            i += 1

        # Horizontal rule
        elif line.strip().startswith('---'):
            doc.add_paragraph()
            i += 1

        # Lists (unordered)
        elif line.strip().startswith('- '):
            list_text = line.replace('- ', '').strip()
            p = doc.add_paragraph(list_text, style='List Bullet')
            i += 1

        # Lists (ordered)
        elif re.match(r'^\d+\. ', line.strip()):
            list_text = re.sub(r'^\d+\. ', '', line.strip())
            p = doc.add_paragraph(list_text, style='List Number')
            i += 1

        # Bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    p.add_run(part)
                else:
                    run = p.add_run(part)
                    run.bold = True
            i += 1

        # Inline code
        elif '`' in line:
            p = doc.add_paragraph()
            parts = line.split('`')
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    p.add_run(part)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            i += 1

        # Tables
        elif '|' in line and i > 0:
            # Check if this is a table header
            if '|' in lines[i]:
                table_rows = []
                j = i

                # Collect table rows
                while j < len(lines) and '|' in lines[j]:
                    row = [cell.strip() for cell in lines[j].split('|')[1:-1]]
                    table_rows.append(row)
                    j += 1

                if len(table_rows) > 0:
                    # Create table
                    num_cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Light Grid Accent 1'

                    # Fill table
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            # Skip separator row (contains ---)
                            if '---' not in cell_data:
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data

                                # Format header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                    i = j

        # Empty lines
        elif line.strip() == '':
            doc.add_paragraph()
            i += 1

        # Regular paragraphs
        else:
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.line_spacing = 1.15
            i += 1

    return doc


def main():
    print("=" * 70)
    print("📝 CONVERT MARKDOWN BÁOCÁO TO DOCX")
    print("=" * 70)

    # Check if markdown file exists
    if not MD_FILE.exists():
        print(f"❌ File not found: {MD_FILE}")
        return

    print(f"\n📖 Reading: {MD_FILE.name}")

    # Read markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()

    print(f"✓ Read {len(md_content):,} characters")

    # Convert
    print("\n🔄 Converting to DOCX...")
    doc = parse_and_convert(md_content)

    # Add metadata
    core_props = doc.core_properties
    core_props.title = "Báo Cáo Thực Tập Tốt Nghiệp - Sàn Tím Vi En"
    core_props.author = "Nguyễn Văn Sang"
    core_props.subject = "Kiểm Thử Tự Động Với Playwright"
    core_props.created = datetime.now()

    # Save
    doc.save(DOCX_FILE)

    print(f"✅ Saved: {DOCX_FILE.name}")
    print(f"📊 File size: {DOCX_FILE.stat().st_size:,} bytes")
    print(f"   ({DOCX_FILE.stat().st_size / 1024 / 1024:.2f} MB)")

    print("\n" + "=" * 70)
    print("✨ Conversion complete!")
    print("=" * 70)


if __name__ == '__main__':
    main()
